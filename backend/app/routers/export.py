from fastapi import APIRouter, HTTPException, Response
from pydantic import BaseModel
from typing import Optional
from fpdf import FPDF
from docx import Document
import docx.shared
import io
import re
import pandas as pd
from ..services.gemini_service import gemini_service
import json

from ..auth import get_current_user
from ..models import User
from fastapi import Depends

router = APIRouter()

@router.get("/test")
async def test_export_env():
    try:
        import fpdf
        import docx
        return {
            "status": "ok", 
            "message": "Environment is healthy",
            "fpdf_version": getattr(fpdf, '__version__', 'unknown'),
            "docx_version": getattr(docx, '__version__', 'unknown')
        }
    except Exception as e:
        return {"status": "error", "detail": f"Dependency missing: {str(e)}"}

class ExportRequest(BaseModel):
    content: str
    filename: Optional[str] = "document"

def _load_unicode_font(pdf):
    """Try to load a Unicode TTF font. Returns font family name or None."""
    import os
    font_candidates = [
        # Linux / Railway (Docker)
        ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
         '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
         '/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf',
         '/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf'),
        # macOS
        ('/System/Library/Fonts/Supplemental/Arial.ttf',
         '/System/Library/Fonts/Supplemental/Arial Bold.ttf',
         '/System/Library/Fonts/Supplemental/Arial Italic.ttf',
         '/System/Library/Fonts/Supplemental/Arial Bold Italic.ttf'),
    ]
    for regular, bold, italic, bold_italic in font_candidates:
        if os.path.exists(regular):
            try:
                pdf.add_font('UniFont', '', regular)
                pdf.add_font('UniFont', 'B', bold if os.path.exists(bold) else regular)
                pdf.add_font('UniFont', 'I', italic if os.path.exists(italic) else regular)
                pdf.add_font('UniFont', 'BI', bold_italic if os.path.exists(bold_italic) else regular)
                print(f"DEBUG: Loaded Unicode font from {regular}")
                return 'UniFont'
            except Exception as e:
                print(f"DEBUG: Font load failed for {regular}: {e}")
                continue
    return None

def _sanitize_for_latin1(text):
    """Replace Unicode chars not in latin-1 with ASCII equivalents."""
    replacements = {
        '\u2019': "'", '\u2018': "'",
        '\u201c': '"', '\u201d': '"',
        '\u2013': '-', '\u2014': '--',
        '\u2026': '...', '\u00a0': ' ',
        '\u2022': '-', '\u2192': '->',
        '\u2610': '[ ]', '\u2611': '[x]', '\u2612': '[x]',
        '\u25cf': '-', '\u25cb': 'o',
        '\u2713': 'v', '\u2717': 'x',
        '\u20ac': 'EUR',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode('latin-1', 'replace').decode('latin-1')

def md_to_pdf(md_text):
    """
    Converts Markdown to PDF using FPDF2 + markdown library for proper HTML rendering.
    Supports: bold, italic, headers, tables, lists, landscape pages for grids.
    """
    print(f"DEBUG: Starting PDF generation... Content len: {len(md_text)}")
    try:
        from fpdf import FPDF
        import markdown
        import os

        class PDF(FPDF):
            def footer(self):
                self.set_y(-15)
                self.set_font('Helvetica', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}', 0, align='C')

        pdf = PDF()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Load Unicode font for proper French character support
        font_name = _load_unicode_font(pdf)
        if not font_name:
            print("DEBUG: No Unicode font found, sanitizing text for latin-1")
            md_text = _sanitize_for_latin1(md_text)
            font_name = 'Helvetica'

        # Split content into sections by "---" horizontal rules
        sections = re.split(r'\n---+\n', md_text)

        for idx, section in enumerate(sections):
            section = section.strip()
            if not section:
                continue

            # Check if this section needs landscape (grille d'aide)
            is_landscape = bool(re.search(r"GRILLE D.AIDE", section, re.IGNORECASE))

            if idx == 0:
                pdf.add_page()
            else:
                pdf.add_page(orientation='L' if is_landscape else 'P')

            # Convert markdown to HTML with table support
            html = markdown.markdown(section, extensions=['tables'])

            # FPDF2 limitation: no nested HTML tags inside <td>/<th>
            # Strip formatting tags and convert <br> to newlines
            def clean_cell(match):
                content = match.group(2)
                content = re.sub(r'<br\s*/?>', '\n', content)
                content = re.sub(r'</?(?:strong|em|b|i|code|p)>', '', content)
                return match.group(1) + content + match.group(3)
            html = re.sub(r'(<t[dh][^>]*>)(.*?)(</t[dh]>)', clean_cell, html, flags=re.DOTALL)

            # Style tables for PDF rendering
            html = html.replace('<table>', '<table border="1" cellpadding="5" width="100%">')
            html = html.replace('<th>', '<th bgcolor="#DDDDDD">')

            # Render HTML
            pdf.set_font(font_name, size=11)
            pdf.write_html(html)

        output = bytes(pdf.output())
        print(f"DEBUG: PDF generation success, bytes: {len(output)}")
        return output

    except Exception as e:
        print(f"❌ PDF generation error: {e}")
        import traceback
        traceback.print_exc()
        # Fallback: plain text PDF
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        safe_text = md_text.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, safe_text)
        return bytes(pdf.output())

def md_to_docx(md_text):
    """
    Converts Markdown to DOCX using pure python-docx with improved Table parsing and Layout control.
    Supports Landscape mode for Grids.
    """
    print(f"DEBUG: Starting DOCX generation... Content len: {len(md_text)}")
    try:
        doc = Document()
        from docx.enum.section import WD_SECTION, WD_ORIENT
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Helper to set narrow margins
        def set_narrow_margins(section):
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Initial Section
        set_narrow_margins(doc.sections[0])

        lines = md_text.split('\n')
        iterator = iter(lines)
        
        in_table = False
        table_lines = []

        def flush_table(lines_to_flush):
            if not lines_to_flush: return
            
            # Parse header
            header_row = lines_to_flush[0].strip().split('|')[1:-1]
            header_row = [h.strip() for h in header_row]
            
            # Create Table
            table = doc.add_table(rows=1, cols=len(header_row))
            table.style = 'Table Grid'
            
            # Fill Header
            hdr_cells = table.rows[0].cells
            for i, h_text in enumerate(header_row):
                if i < len(hdr_cells):
                    hdr_cells[i].text = h_text
                    # Make header bold
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            
            # Fill Rows
            for line in lines_to_flush[1:]:
                # skip separator
                if '---' in line: continue
                
                row_data = line.strip().split('|')[1:-1]
                row_data = [d.strip() for d in row_data]
                
                # Check mismatch cols
                if len(row_data) != len(header_row):
                     # Simple logic: pad or truncate
                     if len(row_data) < len(header_row):
                         row_data += [''] * (len(header_row) - len(row_data))
                     else:
                         row_data = row_data[:len(header_row)]
                
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(row_data):
                    # Handle breaks <br>
                    clean_text = cell_text.replace('<br>', '\n').replace('<br/>', '\n')
                    row_cells[i].text = clean_text

        for line in iterator:
            stripped = line.strip()
            
            # Table Detection
            if stripped.startswith('|'):
                in_table = True
                table_lines.append(stripped)
                continue
            else:
                if in_table:
                    flush_table(table_lines)
                    table_lines = []
                    in_table = False
                    doc.add_paragraph() # Spacer

            # 1. Section/Page Breaks for Grids
            # Detect "PAGE 2 : GRILLE" or similar keywords
            if stripped.upper().startswith('# PAGE 2') or stripped.upper().startswith('## PAGE 2') or "GRILLE D'AIDE" in stripped.upper():
                # Add Section Break
                new_section = doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.orientation = WD_ORIENT.LANDSCAPE
                new_section.page_width = Inches(11.69) # A4 Landscape width
                new_section.page_height = Inches(8.27) # A4 Landscape height
                set_narrow_margins(new_section)
                
            # 2. Headers
            if stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=0)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=1)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=2)
            
            # 3. Lists
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', stripped):
                # Remove number
                parts = stripped.split('.', 1)
                content = parts[1].strip() if len(parts) > 1 else stripped
                doc.add_paragraph(content, style='List Number')
            
            # 4. Standard Text
            elif stripped:
                p = doc.add_paragraph(stripped)
                # Keep With Next check for titles (simple heuristic)
                if len(stripped) < 50 and stripped.endswith(':'):
                    p.paragraph_format.keep_with_next = True

        # Flush trailing table
        if in_table:
            flush_table(table_lines)

        result = io.BytesIO()
        doc.save(result)
        docx_bytes = result.getvalue()
        print(f"DEBUG: DOCX generation success, bytes: {len(docx_bytes)}")
        return docx_bytes

    except Exception as e:
        print(f"❌ Pure Python DOCX Error: {e}")
        # Fallback to absolute basic
        doc = Document()
        doc.add_paragraph(f"Error generating formatted doc: {e}\n\nRaw Content:\n{md_text}")
        result = io.BytesIO()
        doc.save(result)
        return result.getvalue()

@router.post("/pdf")
async def export_pdf(request: ExportRequest, current_user: User = Depends(get_current_user)):
    print(f"DEBUG: Export PDF Request Received. Content len: {len(request.content)}")
    try:
        import unicodedata
        safe_filename = unicodedata.normalize('NFKD', request.filename).encode('ascii', 'ignore').decode('ascii')
        safe_filename = re.sub(r'[^a-zA-Z0-9_\-]', '_', safe_filename)
        
        pdf_bytes = md_to_pdf(request.content)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={safe_filename}.pdf",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        print(f"❌ PDF Export Error CRASH: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/docx")
async def export_docx(request: ExportRequest, current_user: User = Depends(get_current_user)):
    print(f"DEBUG: Export DOCX Request Received. Content len: {len(request.content)}")
    try:
        import unicodedata
        safe_filename = unicodedata.normalize('NFKD', request.filename).encode('ascii', 'ignore').decode('ascii')
        safe_filename = re.sub(r'[^a-zA-Z0-9_\-]', '_', safe_filename)

        docx_bytes = md_to_docx(request.content)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={safe_filename}.docx",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        print(f"❌ DOCX Export Error CRASH: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

# --- Specialized Quiz Exports ---

@router.post("/quiz/gift")
async def export_gift(request: ExportRequest, current_user: User = Depends(get_current_user)):
    """
    Transforms a Markdown quiz into Moodle GIFT format using Gemini.
    """
    try:
        prompt = f"""Tu es un expert en Moodle (format GIFT). Transforme ce quiz Markdown en format GIFT (.txt) valide.
        Règles CRITIQUES :
        1. Chaque question DOIT être suivie d'une ligne vide.
        2. Format QCM : La question {{=RéponseCorrecte ~MauvaiseRéponse1 ~MauvaiseRéponse2}}
        3. Pas de titres de parties (ex: ### Partie 1), juste les questions.
        4. Échappe les caractères spéciaux si nécessaire (ex: ~ , = , # , {{ , }}).
        5. Ne réponds QUE avec le texte GIFT pur. Pas de markdown, pas de ```.

        Quiz à transformer :
        {request.content}
        """
        
        response = gemini_service.client.models.generate_content(
            model=gemini_service.model_name,
            contents=prompt
        )
        gift_content = response.text.strip()
        
        # Clean potential AI noise
        gift_content = re.sub(r'^```[\w]*\n?', '', gift_content)
        gift_content = re.sub(r'\n?```$', '', gift_content)
        
        return Response(
            content=gift_content.encode('utf-8'),
            media_type="text/plain",
            headers={
                "Content-Disposition": f"attachment; filename={request.filename}_moodle.txt",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        print(f"❌ GIFT Export Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/quiz/wooclap")
async def export_wooclap(request: ExportRequest, current_user: User = Depends(get_current_user)):
    """
    Transforms a Markdown quiz into an Excel file for Wooclap using the specific template.
    """
    try:
        prompt = f"""Transforme ce quiz Markdown en un JSON structuré pour Excel (Wooclap). 
        Utilise EXACTEMENT cette structure de colonnes pour chaque objet :
        - "Type": "MCQ" (toujours MCQ pour l'instant)
        - "Title": [Le texte de la question]
        - "Correct": [L'index de la bonne réponse : 1, 2, 3 ou 4]
        - "Choice 1": [Option A]
        - "Choice 2": [Option B]
        - "Choice 3": [Option C]
        - "Choice 4": [Option D]

        Règles :
        - Ne renvoie que le JSON (liste d'objets). Pas de texte explicatif.
        - Si une question n'est pas un QCM, ignore-la.

        Quiz :
        {request.content}
        """
        
        response = gemini_service.client.models.generate_content(
            model=gemini_service.model_name,
            contents=prompt
        )
        
        json_str = response.text.strip()
        json_str = re.sub(r'^```json\n?', '', json_str)
        json_str = re.sub(r'\n?```$', '', json_str)
        
        data = json.loads(json_str)
        df = pd.DataFrame(data)
        
        # Ensure column order matches user image exactly
        expected_cols = ["Type", "Title", "Correct", "Choice 1", "Choice 2", "Choice 3", "Choice 4"]
        # Filter and reorder columns that exist
        current_cols = [c for c in expected_cols if c in df.columns]
        df = df[current_cols]
        
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Wooclap')
            
        return Response(
            content=output.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": f"attachment; filename={request.filename}_wooclap.xlsx",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        print(f"❌ Wooclap Export Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/quiz/google")
async def export_google(request: ExportRequest, current_user: User = Depends(get_current_user)):
    """
    Transforms a Markdown quiz into a CSV for Google Forms imports.
    """
    try:
        prompt = f"""Transforme ce quiz Markdown en format CSV (séparateur virgule) prêt pour Google Forms.
        Colonnes : Question, Option 1, Option 2, Option 3, Option 4, Correct Answer
        Ne réponds QUE avec le CSV brut. Pas de blabla.

        Quiz Markdown :
        {request.content}
        """
        
        response = gemini_service.client.models.generate_content(
            model=gemini_service.model_name,
            contents=prompt
        )
        
        csv_content = response.text.strip()
        csv_content = re.sub(r'^```csv\n?', '', csv_content)
        csv_content = re.sub(r'\n?```$', '', csv_content)
        
        return Response(
            content=csv_content.encode('utf-8'),
            media_type="text/csv",
            headers={
                "Content-Disposition": f"attachment; filename={request.filename}_google.csv",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        print(f"❌ Google Forms Export Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
