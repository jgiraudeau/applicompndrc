from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from sqlalchemy.orm import Session
from ..services.gemini_service import gemini_service
from ..database import get_db
from ..auth import get_current_active_user
from ..models import User, SavedDocument
from pydantic import BaseModel
from typing import List
import shutil
import os
import tempfile
import docx

router = APIRouter()

# --- Schemas ---
class SaveDocRequest(BaseModel):
    title: str
    content: str
    document_type: str

class SavedDocResponse(BaseModel):
    id: str
    title: str
    content: str
    document_type: str
    created_at: str

    class Config:
        from_attributes = True

# --- Endpoints ---

@router.post("/save")
async def save_document_endpoint(
    doc: SaveDocRequest, 
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    try:
        new_doc = SavedDocument(
            user_id=current_user.id,
            title=doc.title,
            content=doc.content,
            document_type=doc.document_type
        )
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)
        return {"status": "success", "id": new_doc.id}
    except Exception as e:
        print(f"Save error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/list")
async def list_documents(
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    # Returns lightweight list (without full content if needed, but for now full is fine)
    # To optimize: db.query(SavedDocument.id, SavedDocument.title...)
    docs = db.query(SavedDocument).filter(SavedDocument.user_id == current_user.id).order_by(SavedDocument.created_at.desc()).all()
    # Normalize created_at to string
    return docs

@router.get("/{doc_id}")
async def get_document(
    doc_id: str,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    doc = db.query(SavedDocument).filter(SavedDocument.id == doc_id, SavedDocument.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

@router.delete("/{doc_id}")
async def delete_document(
    doc_id: str,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    doc = db.query(SavedDocument).filter(SavedDocument.id == doc_id, SavedDocument.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    db.delete(doc)
    db.commit()
    return {"status": "deleted"}

@router.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """
    Uploads a file (PDF, etc.) to Gemini and returns the file handle.
    Converts DOCX to TXT before upload.
    """
    try:
        # Save temp file
        original_suffix = f"_{file.filename}"
        with tempfile.NamedTemporaryFile(delete=False, suffix=original_suffix) as tmp:
            shutil.copyfileobj(file.file, tmp)
            tmp_path = tmp.name

        # Determine mime type
        mime_type = file.content_type or "application/pdf"
        final_path = tmp_path
        final_mime = mime_type

        # Handle DOCX -> TXT conversion
        if "wordprocessingml" in mime_type or file.filename.endswith(".docx"):
            print(f"🔄 Converting {file.filename} from DOCX to TXT...")
            try:
                doc = docx.Document(tmp_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                text_content = "\n".join(full_text)
                
                # Setup new txt file
                txt_path = tmp_path + ".txt"
                with open(txt_path, "w", encoding="utf-8") as f:
                    f.write(text_content)
                
                # Switch to new file for upload
                final_path = txt_path
                final_mime = "text/plain"
            except Exception as e:
                print(f"⚠️ DOCX Conversion failed: {e}")
                # Fallback to original attempts
        
        # Upload to Gemini
        gemini_file = gemini_service.upload_file_to_gemini(final_path, mime_type=final_mime)
        
        # Cleanup temp file(s)
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        if final_path != tmp_path and os.path.exists(final_path):
            os.unlink(final_path)

        return {
            "status": "success",
            "filename": file.filename,
            "gemini_file_name": gemini_file.name,
            "gemini_file_uri": gemini_file.uri
        }

    except Exception as e:
        print(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
